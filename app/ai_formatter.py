# C:\Users\Pavlu4ini\PycharmProjects\TranscriberAssistant\app\ai_formatter.py

from docx import Document

def format_document(file_info, transcribed_text):
    """
    Оформляем документ на основе текста транскрипции и информации о ролях участников.
    """
    doc = Document()
    doc.add_heading('Транскрипция', 0)

    # Добавляем информацию о ролях
    for participant, role in zip(file_info["participants"], file_info["roles"]):
        doc.add_paragraph(f'{participant} ({role})')

    # Добавляем сам текст
    doc.add_paragraph(transcribed_text)

    # Сохраняем документ
    doc_path = "downloads/transcription.docx"
    doc.save(doc_path)

    return doc_path
